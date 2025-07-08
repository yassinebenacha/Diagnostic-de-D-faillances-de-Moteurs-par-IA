import openai
import os
from typing import Dict, Any, Optional
from datetime import datetime
import pdfkit
from docx import Document

# Templates de prompts pour différents types de défaillances
PROMPT_TEMPLATES = {
    'fr': {
        'default': """
Vous êtes un expert en maintenance industrielle. Rédigez un rapport de diagnostic pour un moteur électrique à partir des résultats suivants :
- Prédiction : {prediction}
- Confiance : {confidence:.2%}
- Dérive détectée : {drift}
- Caractéristiques extraites : {features}
- Recommandations de maintenance :
{recommendations}
""",
        'bearing_fault': """
Un défaut de roulement a été détecté. Résumez les indicateurs clés et proposez des actions correctives adaptées à la criticité.
Résultats :
- Prédiction : {prediction}
- Confiance : {confidence:.2%}
- Caractéristiques : {features}
- Dérive : {drift}
- Conseils : {recommendations}
"""
    },
    'en': {
        'default': """
You are an industrial maintenance expert. Write a diagnostic report for an electric motor based on the following results:
- Prediction: {prediction}
- Confidence: {confidence:.2%}
- Drift detected: {drift}
- Extracted features: {features}
- Maintenance recommendations:
{recommendations}
""",
        'bearing_fault': """
A bearing fault was detected. Summarize the key indicators and propose corrective actions according to criticality.
Results:
- Prediction: {prediction}
- Confidence: {confidence:.2%}
- Features: {features}
- Drift: {drift}
- Advice: {recommendations}
"""
    }
}

# Recommandations de maintenance par défaut
DEFAULT_RECOMMENDATIONS = {
    'fr': {
        'Normal': "Aucune action requise. Continuer la surveillance régulière.",
        'Défaillance': "Planifier une inspection approfondie et envisager le remplacement des composants défectueux."
    },
    'en': {
        'Normal': "No action required. Continue regular monitoring.",
        'Fault': "Schedule a thorough inspection and consider replacing faulty components."
    }
}

class DiagnosticReportGenerator:
    """
    Générateur de rapports de diagnostic automatisé via LLM (OpenAI API).
    """
    def __init__(self, openai_api_key: str, language: str = 'fr'):
        openai.api_key = openai_api_key
        self.language = language if language in PROMPT_TEMPLATES else 'fr'

    def format_features(self, features: Dict[str, Any]) -> str:
        return '\n'.join([f"- {k}: {v:.4f}" if isinstance(v, float) else f"- {k}: {v}" for k, v in features.items()])

    def get_recommendations(self, prediction: str) -> str:
        return DEFAULT_RECOMMENDATIONS[self.language].get(prediction, "Consulter un expert pour analyse approfondie.")

    def build_prompt(self, result: Dict[str, Any], fault_type: Optional[str] = None) -> str:
        template = PROMPT_TEMPLATES[self.language].get(fault_type, PROMPT_TEMPLATES[self.language]['default'])
        return template.format(
            prediction=result.get('prediction'),
            confidence=result.get('confidence', 0),
            drift=result.get('drift', False),
            features=self.format_features(result.get('features', {})),
            recommendations=self.get_recommendations(result.get('prediction'))
        )

    def generate_report(self, result: Dict[str, Any], fault_type: Optional[str] = None, quality_check: bool = True) -> str:
        prompt = self.build_prompt(result, fault_type)
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "Vous êtes un expert en diagnostic industriel."},
                      {"role": "user", "content": prompt}]
        )
        report = response['choices'][0]['message']['content']
        if quality_check:
            if not self.validate_report(report):
                raise ValueError("Le rapport généré ne satisfait pas les critères de qualité.")
        return report

    def validate_report(self, report: str) -> bool:
        # Simple check: longueur minimale et présence de sections clés
        return len(report) > 200 and ("recommandation" in report.lower() or "recommendation" in report.lower())

    def export_report(self, report: str, output_path: str, fmt: str = 'html'):
        if fmt == 'html':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"<html><body><h1>Rapport de Diagnostic</h1><pre>{report}</pre></body></html>")
        elif fmt == 'pdf':
            html_path = output_path.replace('.pdf', '.html')
            self.export_report(report, html_path, 'html')
            pdfkit.from_file(html_path, output_path)
        elif fmt == 'docx':
            doc = Document()
            doc.add_heading('Rapport de Diagnostic', 0)
            for line in report.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_path)
        else:
            raise ValueError(f"Format non supporté: {fmt}")

# Exemples de prompts optimisés et templates de rapports
EXAMPLE_PROMPT_FR = PROMPT_TEMPLATES['fr']['default'].format(
    prediction="Défaillance",
    confidence=0.92,
    drift=True,
    features="- rms: 0.1234\n- kurtosis: 3.45",
    recommendations=DEFAULT_RECOMMENDATIONS['fr']['Défaillance']
)

EXAMPLE_PROMPT_EN = PROMPT_TEMPLATES['en']['default'].format(
    prediction="Fault",
    confidence=0.92,
    drift=True,
    features="- rms: 0.1234\n- kurtosis: 3.45",
    recommendations=DEFAULT_RECOMMENDATIONS['en']['Fault']
)

EXAMPLE_REPORT_TEMPLATE = """
<html><body><h1>Rapport de Diagnostic</h1><pre>{report}</pre></body></html>
""" 